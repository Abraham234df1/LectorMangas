from pathlib import Path
import sys

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\manue\Escritorio\Programacion\MangaReadV1(1)\MangaReadV1")
ASSETS = ROOT / "tmp" / "document_assets"
OUTPUT = ROOT / "output" / "MangaReadV1_Entrega_Completa.docx"
SKILL_ROOT = Path(r"C:\Users\manue\.codex\plugins\cache\openai-primary-runtime\documents\26.715.12143\skills\documents")
sys.path.insert(0, str(SKILL_ROOT / "scripts"))
from table_geometry import apply_table_geometry  # noqa: E402


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "0B2545"
INK = "202832"
MUTED = "667085"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
GREEN = "1F6B4F"
GREEN_FILL = "E8F5EF"
AMBER = "7A5A00"
AMBER_FILL = "FFF4D6"
RED = "9B1C1C"
RED_FILL = "FDECEC"
WHITE = "FFFFFF"
BORDER = "C7D0DB"


def rgb(value):
    return RGBColor.from_string(value)


def set_run_font(run, name="Calibri", size=11, color=INK, bold=False, italic=False):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    run.bold = bold
    run.italic = italic


def set_cell_fill(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_border(cell, color=BORDER, size="5"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_keep_with_next(paragraph, value=True):
    paragraph.paragraph_format.keep_with_next = value


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instruction, separate, text, end):
        run._r.append(element)
    set_run_font(run, size=9, color=MUTED)


def add_bullet(doc, text, level=0):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.375 + level * 0.25)
    paragraph.paragraph_format.first_line_indent = Inches(-0.188)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    run = paragraph.add_run(text)
    set_run_font(run)
    return paragraph


def add_number(doc, text):
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.paragraph_format.left_indent = Inches(0.375)
    paragraph.paragraph_format.first_line_indent = Inches(-0.188)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    run = paragraph.add_run(text)
    set_run_font(run)
    return paragraph


def add_body(doc, text, bold_lead=None):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.25
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        set_run_font(lead, bold=True, color=NAVY)
        rest = paragraph.add_run(text[len(bold_lead):])
        set_run_font(rest)
    else:
        run = paragraph.add_run(text)
        set_run_font(run)
    return paragraph


def add_code_block(doc, code):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.15)
    paragraph.paragraph_format.right_indent = Inches(0.15)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.0
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), LIGHT_GRAY)
    p_pr.append(shd)
    for idx, line in enumerate(code.splitlines()):
        if idx:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        set_run_font(run, name="Consolas", size=8.5, color=INK)
    return paragraph


def add_callout(doc, label, text, tone="info"):
    palette = {
        "info": (LIGHT_BLUE, NAVY),
        "ok": (GREEN_FILL, GREEN),
        "warning": (AMBER_FILL, AMBER),
        "risk": (RED_FILL, RED),
    }
    fill, color = palette[tone]
    table = doc.add_table(rows=1, cols=1)
    # A one-row callout is still represented as a table in Word. Mark the row
    # explicitly so screen readers and the document audit do not encounter an
    # unlabelled table structure.
    set_repeat_header(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_fill(cell, fill)
    set_cell_border(cell, color=color, size="7")
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.15
    label_run = paragraph.add_run(f"{label}: ")
    set_run_font(label_run, bold=True, color=color)
    text_run = paragraph.add_run(text)
    set_run_font(text_run, color=INK)
    apply_table_geometry(table, [9360], indent_dxa=120)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_table(doc, headers, rows, widths, status_column=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header = table.rows[0]
    set_repeat_header(header)
    for idx, text in enumerate(headers):
        cell = header.cells[idx]
        set_cell_fill(cell, LIGHT_BLUE)
        set_cell_border(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(str(text))
        set_run_font(run, size=9.5, color=NAVY, bold=True)
    for row_data in rows:
        row = table.add_row()
        for idx, value in enumerate(row_data):
            cell = row.cells[idx]
            set_cell_border(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.1
            if idx == status_column:
                normalized = str(value).lower()
                if "aprob" in normalized or "cumple" in normalized or "listo" in normalized:
                    fill, color = GREEN_FILL, GREEN
                elif "pend" in normalized or "bloque" in normalized:
                    fill, color = AMBER_FILL, AMBER
                elif "fall" in normalized:
                    fill, color = RED_FILL, RED
                else:
                    fill, color = LIGHT_GRAY, NAVY
                set_cell_fill(cell, fill)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run(str(value))
                set_run_font(run, size=9, color=color, bold=True)
            else:
                run = paragraph.add_run(str(value))
                set_run_font(run, size=9.2, color=INK)
    apply_table_geometry(table, widths, indent_dxa=120)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)
    return table


def add_caption(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(10)
    paragraph.paragraph_format.keep_with_next = False
    run = paragraph.add_run(text)
    set_run_font(run, size=9, color=MUTED, italic=True)
    return paragraph


def set_picture_alt(inline_shape, description):
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("descr", description)
    doc_pr.set("title", description)


def add_figure(doc, filename, caption, width=5.85):
    path = ASSETS / filename
    if not path.exists():
        raise FileNotFoundError(path)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run()
    shape = run.add_picture(str(path), width=Inches(width))
    set_picture_alt(shape, caption)
    add_caption(doc, caption)


def configure_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in tokens.items():
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_style in ("List Bullet", "List Number"):
        style = doc.styles[list_style]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def configure_header_footer(doc):
    section = doc.sections[0]
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run("MANGAREADV1  |  DOCUMENTO DE ENTREGA")
    set_run_font(run, size=8.5, color=MUTED, bold=True)

    footer = section.footer
    table = footer.add_table(rows=1, cols=2, width=Inches(6.5))
    set_repeat_header(table.rows[0])
    left = table.cell(0, 0).paragraphs[0]
    left.paragraph_format.space_after = Pt(0)
    left_run = left.add_run("Actividad integradora - 22 de julio de 2026")
    set_run_font(left_run, size=8.5, color=MUTED)
    right = table.cell(0, 1).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right.paragraph_format.space_after = Pt(0)
    prefix = right.add_run("Página ")
    set_run_font(prefix, size=8.5, color=MUTED)
    add_page_field(right)
    apply_table_geometry(table, [6800, 2560], indent_dxa=0, cell_margins_dxa={"top": 0, "bottom": 0, "start": 0, "end": 0})
    table._tbl.tblPr.remove(table._tbl.tblPr.find(qn("w:tblBorders"))) if table._tbl.tblPr.find(qn("w:tblBorders")) is not None else None


def add_cover(doc):
    for _ in range(3):
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(16)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(18)
    run = kicker.add_run("ACTIVIDAD INTEGRADORA")
    set_run_font(run, size=11, color=BLUE, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run("MangaReadV1")
    set_run_font(run, size=31, color=NAVY, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(6)
    run = subtitle.add_run("Documento de entrega, evidencias y pruebas")
    set_run_font(run, size=16, color=DARK_BLUE)

    detail = doc.add_paragraph()
    detail.alignment = WD_ALIGN_PARAGRAPH.CENTER
    detail.paragraph_format.space_after = Pt(34)
    run = detail.add_run("Biblioteca digital de mangas, cómics y novelas ligeras en PDF")
    set_run_font(run, size=11, color=MUTED, italic=True)

    add_table(
        doc,
        ["Indicador", "Resultado", "Cobertura"],
        [
            ("Catálogo", "13 títulos", "Biblioteca pública"),
            ("Lecturas", "52 disponibles", "Manga y novela ligera"),
            ("Lector", "PDF.js + PageFlip", "LTR, RTL y responsive"),
        ],
        [2500, 2300, 4560],
    )

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(26)
    note.paragraph_format.space_after = Pt(4)
    run = note.add_run("Equipo de desarrollo MangaReadV1")
    set_run_font(run, size=11, color=NAVY, bold=True)
    date = doc.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date.add_run("22 de julio de 2026")
    set_run_font(run, size=10.5, color=MUTED)
    doc.add_page_break()


def add_section_title(doc, text):
    paragraph = doc.add_heading(text, level=1)
    set_keep_with_next(paragraph)
    return paragraph


def build_document():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)
    configure_header_footer(doc)
    add_cover(doc)

    add_section_title(doc, "1. Resumen ejecutivo")
    add_body(doc, "MangaReadV1 es una aplicación web modular construida con HTML, CSS y JavaScript. Permite consultar una biblioteca pública, buscar títulos, revisar detalles, abrir PDFs en un lector con efecto de cambio de hoja e ingresar a un panel administrativo protegido.")
    add_body(doc, "La versión local está operativa con 13 títulos y 52 lecturas visibles. Incluye 14 bloques de capítulos de Classroom of the Elite y 38 archivos de su novela ligera organizados en Año 1, Año 2 y Año 3. El lector fue comprobado con PDFs reales y navegación responsiva.")
    add_callout(doc, "Estado de Supabase", "La URL y la clave pública están configuradas, pero el proyecto remoto aún no tiene desplegadas todas las tablas, buckets y usuarios requeridos. Por ello, la aplicación usa un respaldo local y bloquea las escrituras remotas para no simular persistencia.", "warning")

    doc.add_heading("Alcance solicitado", level=2)
    add_figure(doc, "00-requisitos-entregables.png", "Figura 1. Lista de entregables proporcionada para la actividad.", width=4.45)

    doc.add_heading("Mapa del documento", level=2)
    for item in [
        "Matriz de cumplimiento y estado real de cada entregable.",
        "Inventario y estructura del código fuente.",
        "Evidencias de Supabase y pasos pendientes de despliegue.",
        "Capturas de biblioteca, detalle, login, formularios y lector.",
        "Documento de pruebas con resultados verificables.",
        "Explicación de conexión, seguridad, validaciones y dificultades.",
    ]:
        add_bullet(doc, item)

    add_section_title(doc, "2. Matriz de entregables")
    deliverables = [
        ("Código fuente completo", "Cumple", "205 archivos del proyecto; 6 HTML, 27 JS, 7 CSS y 4 SQL."),
        ("Captura de tablas Supabase", "Pendiente remoto", "schema.sql prepara admins, mangas, volumes y chapter_marks; falta ejecutar en el panel."),
        ("Captura de buckets", "Pendiente remoto", "policies.sql prepara covers y pdfs; la API pública devuelve una lista vacía."),
        ("Usuario administrador Supabase", "Pendiente remoto", "El acceso local admin@gmail.com funciona; falta crear/verificar el usuario remoto."),
        ("Capturas de la aplicación", "Cumple", "Biblioteca, detalle, login, formulario de manga, formulario de tomo y lector incluidos."),
        ("Documento de pruebas", "Cumple", "Casos locales aprobados y pruebas remotas identificadas como pendientes."),
        ("Documento explicativo", "Cumple", "Incluye estructura, conexión, protección administrativa, validaciones y dificultades."),
    ]
    add_table(doc, ["Entregable", "Estado", "Evidencia"], deliverables, [2550, 1750, 5060], status_column=1)
    add_callout(doc, "Criterio de honestidad", "El informe distingue entre funcionalidad local comprobada, scripts preparados y recursos realmente desplegados en Supabase.", "info")

    add_section_title(doc, "3. Código fuente y estructura del proyecto")
    add_body(doc, "El código está organizado por responsabilidades: páginas HTML, estilos reutilizables, servicios de datos, módulos de interfaz, lector PDF y scripts SQL. La separación facilita mantenimiento, pruebas y sustitución del respaldo local por Supabase cuando el backend quede desplegado.")
    add_table(
        doc,
        ["Componente", "Cantidad", "Responsabilidad"],
        [
            ("HTML", "6", "Biblioteca, detalle, lector, administración, login y registro."),
            ("JavaScript", "27", "Servicios, autenticación, validaciones, interfaz y lector."),
            ("CSS", "7", "Diseño general, componentes, administración, lector y responsive."),
            ("SQL", "4", "Esquema, políticas, catálogo y datos iniciales."),
            ("PDF", "64", "52 lecturas activas más 12 PDFs demo heredados."),
        ],
        [1900, 1200, 6260],
    )
    add_code_block(doc, """MangaReadV1/
|- index.html, manga.html, reader.html
|- login.html, register.html, admin.html
|- css/          estilos, componentes y responsive
|- js/           servicios, UI, autenticación y lector
|- supabase/     schema.sql, policies.sql, catalog.sql
|- output/pdf/   PDFs organizados por colección y año
|- docs/         auditoría, pruebas y documentación""")
    doc.add_heading("Flujo funcional", level=2)
    for step in [
        "index.html carga listMangas(), mezcla Supabase con el catálogo local y dibuja las tarjetas.",
        "manga.html recibe el identificador, carga metadatos y presenta tomos o secciones por año.",
        "reader.html descarga el PDF, lo procesa con PDF.js y lo muestra mediante StPageFlip.",
        "login.html crea una sesión autorizada; guards.js impide entrar al panel sin autorización.",
        "admin.html coordina el CRUD remoto o la edición local de lecturas existentes.",
    ]:
        add_number(doc, step)

    doc.add_page_break()
    add_section_title(doc, "4. Evidencias de Supabase")
    add_body(doc, "La aplicación apunta al proyecto rbpueniwctygjjkdqhny.supabase.co mediante una clave pública. La clave service_role no se incluye en el frontend.")
    add_table(
        doc,
        ["Comprobación pública", "Respuesta", "Interpretación"],
        [
            ("REST /mangas", "HTTP 200 - []", "La conexión responde, pero no hay registros remotos."),
            ("REST /admins", "HTTP 404", "La tabla de administradores no está desplegada."),
            ("Storage /bucket", "HTTP 200 - []", "No hay buckets visibles para la clave pública."),
            ("Dashboard", "Requiere iniciar sesión", "No fue posible capturar Table Editor, Storage ni Auth sin la cuenta propietaria."),
        ],
        [2500, 1850, 5010],
    )
    add_figure(doc, "07-supabase-acceso-pendiente.png", "Figura 2. El panel del proyecto solicita autenticación del propietario; las capturas internas quedan pendientes.", width=6.2)

    doc.add_heading("Scripts preparados", level=2)
    add_bullet(doc, "schema.sql crea admins, mangas, volumes y chapter_marks, con relaciones y restricciones de unicidad.")
    add_bullet(doc, "policies.sql activa RLS, define is_admin(), RPC seguras y buckets covers/pdfs.")
    add_bullet(doc, "catalog.sql contiene registros iniciales de catálogo.")
    add_bullet(doc, "seed.example.sql registra el correo administrador inicial.")
    add_code_block(doc, """-- Orden de despliegue en SQL Editor
1. supabase/schema.sql
2. supabase/policies.sql
3. supabase/catalog.sql
4. supabase/seed.example.sql  -- opcional / ajustar correo""")
    add_callout(doc, "Capturas obligatorias pendientes", "Después de ejecutar los SQL con la cuenta propietaria se deben insertar tres capturas: Table Editor con cuatro tablas, Storage con covers/pdfs y Authentication > Users con el administrador.", "risk")

    doc.add_page_break()
    add_section_title(doc, "5. Evidencias de la aplicación")
    add_body(doc, "Las siguientes capturas fueron tomadas directamente de la ejecución local en http://127.0.0.1:8765.")
    add_figure(doc, "01-biblioteca.png", "Figura 3. Biblioteca pública con 13 títulos, 52 lecturas y acceso a las colecciones.")
    doc.add_page_break()
    add_figure(doc, "02-detalle-novela.png", "Figura 4. Detalle de Classroom of the Elite - Novela ligera, con 38 lecturas organizadas en tres años.")
    doc.add_page_break()
    add_figure(doc, "03-login-administrador.png", "Figura 5. Inicio de sesión separado del registro de cuentas administrativas.")
    doc.add_page_break()
    add_figure(doc, "04-formulario-manga.png", "Figura 6. Panel administrativo y formulario de registro de manga.")
    doc.add_page_break()
    add_figure(doc, "05-formulario-tomo.png", "Figura 7. Formulario de tomo con PDF y marcas de capítulo asociado a la novela ligera.")
    doc.add_page_break()
    add_figure(doc, "06-lector-pdf.png", "Figura 8. Lector PDF real del Año 3, volumen 1, con 232 páginas, zoom y cambio de hoja.")

    doc.add_page_break()
    add_section_title(doc, "6. Documento de pruebas")
    add_body(doc, "Las pruebas se separan por alcance. Aprobado local significa que se observó el resultado en la aplicación local. Pendiente remoto identifica acciones que requieren desplegar el backend o iniciar sesión en el panel propietario.")
    tests = [
        ("P01", "Abrir biblioteca", "13 tarjetas y conteo de 52 lecturas", "Aprobado local"),
        ("P02", "Abrir detalle", "Novela ligera con 38 lecturas y tres años", "Aprobado local"),
        ("P03", "Cargar PDF real", "Año 3, volumen 1 muestra / 232", "Aprobado local"),
        ("P04", "Página siguiente", "El contador avanza 1 a 2 con animación", "Aprobado local"),
        ("P05", "Diseño responsive", "Sin desbordamiento horizontal en vista estrecha", "Aprobado local"),
        ("P06", "Login inicial", "admin@gmail.com abre el panel local", "Aprobado local"),
        ("P07", "Panel administrativo", "Muestra 13 mangas y 52 tomos/lecturas", "Aprobado local"),
        ("P08", "Formulario manga", "Campos, portada y sentido de lectura visibles", "Aprobado local"),
        ("P09", "Formulario tomo", "Título, PDF y marcas vinculados al manga", "Aprobado local"),
        ("P10", "Validación JavaScript", "Módulos principales pasan node --check", "Aprobado local"),
        ("P11", "Conexión REST mangas", "HTTP 200 con arreglo vacío", "Aprobado parcial"),
        ("P12", "Tabla admins", "La API devuelve HTTP 404", "Pendiente remoto"),
        ("P13", "Buckets Storage", "La API devuelve una lista vacía", "Pendiente remoto"),
        ("P14", "CRUD persistente", "Crear, editar y eliminar en Supabase", "Pendiente remoto"),
        ("P15", "Usuario Auth", "Usuario administrador visible en Dashboard", "Pendiente remoto"),
    ]
    add_table(doc, ["ID", "Caso", "Resultado observado", "Estado"], tests, [800, 2600, 4100, 1860], status_column=3)
    add_callout(doc, "Balance", "10 pruebas locales aprobadas, 1 conexión parcial y 4 verificaciones remotas pendientes. No se contabilizan como aprobadas las funciones que dependen del backend aún no desplegado.", "info")

    add_section_title(doc, "7. Conexión con Supabase")
    add_body(doc, "config.js define Project URL, anon public key, correos permitidos y nombres de buckets. supabase-client.js crea una sola instancia del SDK y conserva la sesión del usuario.")
    add_code_block(doc, """client = window.supabase.createClient(config.url, config.anonKey, {
  auth: {
    persistSession: true,
    autoRefreshToken: true,
    detectSessionInUrl: true
  }
});""")
    add_body(doc, "Los servicios manga-service.js, volume-service.js, chapter-service.js y storage-service.js encapsulan consultas, inserciones, actualizaciones, eliminación y transferencia de archivos. La interfaz no escribe directamente en las tablas.")
    add_body(doc, "Para PDFs mayores de 45 MB, storage-service.js divide el archivo en partes de 20 MB. El lector descarga las partes en orden y reconstruye un Uint8Array antes de enviarlo a PDF.js.")

    add_section_title(doc, "8. Protección de operaciones administrativas")
    add_bullet(doc, "guards.js exige una sesión y verifica que el correo pertenezca a admins antes de permitir admin.html.")
    add_bullet(doc, "RLS permite SELECT público, pero INSERT, UPDATE y DELETE requieren public.is_admin().")
    add_bullet(doc, "public.is_admin() compara el correo del JWT con public.admins.")
    add_bullet(doc, "register_admin_email() usa security definer y exige el código de invitación.")
    add_bullet(doc, "La anon key es pública; la service_role nunca se coloca en el navegador.")
    add_bullet(doc, "En modo local, las escrituras remotas quedan deshabilitadas y las ediciones de PDFs locales se guardan en IndexedDB.")
    add_callout(doc, "Importante", "El acceso local admin@gmail.com / admin123 sirve para revisión académica. Antes de publicar, se debe desplegar Supabase y cambiar la contraseña inicial.", "warning")

    add_section_title(doc, "9. Validaciones implementadas")
    validation_rows = [
        ("Campos obligatorios", "Título de manga, autor, título de tomo, correo y contraseña."),
        ("Correo", "Formato usuario@dominio y normalización a minúsculas."),
        ("Contraseña", "Mínimo de 8 caracteres y confirmación coincidente."),
        ("Portada", "Debe ser imagen; máximo 10 MB."),
        ("PDF", "MIME o extensión PDF; máximo 800 MB."),
        ("Duplicados", "Título normalizado y nombre de PDF únicos por manga."),
        ("Marcas", "Capítulo y página enteros mayores que cero."),
        ("Unicidad de marcas", "No se repiten capítulos ni páginas iniciales."),
        ("Asociación", "No se agrega tomo sin seleccionar el manga correcto."),
        ("Servidor", "Constraints SQL y RPC transaccional duplican las validaciones críticas."),
    ]
    add_table(doc, ["Área", "Regla aplicada"], validation_rows, [2500, 6860])

    add_section_title(doc, "10. Dificultades encontradas")
    difficulties = [
        ("Supabase sin esquema remoto", "La clave pública no puede crear tablas ni políticas.", "Se prepararon SQL reproducibles y un modo local honesto que bloquea escrituras remotas."),
        ("PDFs grandes", "Archivos de 45 a 60 MB son sensibles a límites y timeouts.", "Subida por chunks de 20 MB y reconstrucción ordenada."),
        ("Render concurrente", "PDF.js puede cancelar tareas sobre el mismo canvas.", "Cola de promesas, caché y generación de render."),
        ("Sincronización PageFlip", "El contador podía saltar una página en RTL.", "Página lógica pendiente hasta finalizar la animación."),
        ("Diseño responsivo", "El libro podía cortarse en pantallas estrechas.", "Cálculo por ancho y altura, reconstrucción al redimensionar y modo retrato."),
        ("Nombres de novelas", "Los nombres descargados no siempre coincidían con el volumen interno.", "Clasificación por contenido y manifiesto por Año 1, Año 2 y Año 3."),
    ]
    add_table(doc, ["Dificultad", "Causa", "Solución"], difficulties, [2100, 3000, 4260])

    add_section_title(doc, "11. Pasos para cerrar la entrega remota")
    for step in [
        "Iniciar sesión en Supabase con la cuenta propietaria del proyecto.",
        "Ejecutar schema.sql, policies.sql y catalog.sql en ese orden.",
        "Ejecutar seed.example.sql después de confirmar el correo administrador.",
        "Activar Email Provider y registrar las URL locales o publicadas.",
        "Verificar CRUD real y subir una portada y un PDF de prueba.",
        "Capturar Table Editor, Storage y Authentication > Users.",
        "Reemplazar la evidencia pendiente de este informe con esas tres capturas.",
        "Cambiar la contraseña inicial antes de publicar.",
    ]:
        add_number(doc, step)
    add_callout(doc, "Resultado esperado", "Con estos pasos, el catálogo dejará de depender del respaldo local y las operaciones administrativas quedarán persistidas y protegidas por RLS.", "ok")

    add_section_title(doc, "12. Conclusión")
    add_body(doc, "MangaReadV1 cumple en local con la biblioteca, detalle, autenticación de revisión, administración de contenido local y lector PDF 3D responsivo. La organización modular, las validaciones y los scripts SQL cubren los requisitos técnicos de la actividad.")
    add_body(doc, "La única brecha de entrega es el despliegue del esquema, buckets y usuario en el proyecto remoto de Supabase. El documento deja esa brecha identificada y proporciona el procedimiento exacto para cerrarla y obtener las capturas finales del panel.")

    doc.core_properties.title = "MangaReadV1 - Documento de entrega, evidencias y pruebas"
    doc.core_properties.subject = "Actividad integradora MangaReadV1"
    doc.core_properties.author = "Equipo de desarrollo MangaReadV1"
    doc.core_properties.keywords = "MangaReadV1, Supabase, PDF.js, PageFlip, evidencias, pruebas"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
